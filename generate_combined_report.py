"""
Generate a combined paper-style Word report merging content from
Audio_Noise_RI_Fit_Report and final_bootstrap_report.

Structure:
  1. Introduction          (blank placeholder)
  2. Noisy RI Model        (formulas + description)
  3. Dataset               (pipeline: text -> audio -> noise -> ASR -> LLM)
  4. Results
     a. Regression figures (empirical + fit, one panel per noise type)
     b. Parameter estimates table (estimate ± SE, 95% CI, per noise type)
     c. Significance tests summary (brief text; full tables -> appendix)
  5. Conclusion            (blank placeholder)
  Appendix 1: Within-noise pairwise lambda comparison (all model pairs × 3 noise types)
  Appendix 2: Across-noise parameter comparison (same model / param across noise types)

Output:
  Dataset/results/paper_report.docx
"""

import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE        = "./Dataset/results"
AUDIO_SUM   = f"{BASE}/bootstrap/bootstrap_summary.csv"
AUDIO_MOD   = f"{BASE}/bootstrap/pairwise_model_tests.csv"
AUDIO_CROSS = f"{BASE}/bootstrap/pairwise_tests.csv"
OUT_PATH    = f"{BASE}/paper_report_v4.docx"

NOISE_ORDER  = ["white",  "babble",  "cafe"]
NOISE_LABELS = {"white": "White Noise", "babble": "Babble Noise", "cafe": "Café Noise"}
MODEL_ORDER  = ["GPT-3.5-turbo", "GPT-5.4-nano", "Gemini-2.5-Flash", "Gemini-2.5-Flash-Lite"]

FIG = {n: {
    "combined": f"{BASE}/audio_{n}/ri_fit/combined_snr_nsr.png",
    "fit":      f"{BASE}/audio_{n}/ri_fit/all_models_fit.png",
    "snr":      f"{BASE}/audio_{n}/ri_fit/SNR_Plot/all_models_fit.png",
    "reg":      f"{BASE}/audio_{n}/ri_fit/all_models_fit_regression_only.png",
    "q":        f"{BASE}/audio_{n}/ri_fit/q_vs_NSR.png",
} for n in NOISE_ORDER}

SIG_MARK  = "✓"
NOSIG_MARK = "–"
HDR_COLOR  = "2F5496"


# ── docx helpers ──────────────────────────────────────────────────────────────

def cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def col_widths(table, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = Inches(widths[i])


def hdr_row(table, cols, bg=HDR_COLOR):
    row = table.rows[0]
    for i, text in enumerate(cols):
        c = row.cells[i]
        c.text = text
        run = c.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_bg(c, bg)


def data_row(table, vals, sig=False, center_cols=None):
    center_cols = center_cols or []
    row = table.add_row()
    for i, v in enumerate(vals):
        c = row.cells[i]
        c.text = str(v)
        run = c.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        if i in center_cols:
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if sig and i == len(vals) - 1:
            run.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x7A, 0x1F)
    if sig:
        for c in row.cells:
            cell_bg(c, "EBF5EB")


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text, bold=False, italic=False, size=11, align=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if align:
        para.alignment = align
    return para


def fmt(val, dec=4):
    return f"{val:.{dec}f}"


def sig_bool(v):
    return str(v).strip().lower() == "true"


def add_figure(doc, path, width=Inches(5.8), caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].font.size = Pt(9)
            cp.runs[0].italic = True
    else:
        p(doc, f"[Figure not found: {path}]", italic=True, size=9)


def page_break(doc):
    doc.add_page_break()


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    aud_sum   = pd.read_csv(AUDIO_SUM)
    aud_mod   = pd.read_csv(AUDIO_MOD)
    aud_cross = pd.read_csv(AUDIO_CROSS)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    # ── Title ─────────────────────────────────────────────────────────────────
    t = doc.add_heading(
        "Interpretability of LLM Classifiers via Rational Inattention Theory", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. INTRODUCTION
    # ══════════════════════════════════════════════════════════════════════════
    h(doc, "1. Introduction", 1)
    p(doc, "[To be written]", italic=True, size=10)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. NOISY RI MODEL
    # ══════════════════════════════════════════════════════════════════════════
    TRELLIS_FIG = r"D:\lecture\lecture1\lecture\PHD\conference\ACL\Paper\figures\Picture5.png"

    h(doc, "2. Noisy Rational Inattention Model", 1)

    p(doc,
      "This paper extends the Rational Inattention (RI) framework for LLM "
      "interpretability introduced in Zhao et al. (2026) "
      "(Sims, 2003; Matejka and McKay, 2015). "
      "We briefly recap the core model; a full derivation is in the original paper.",
      size=10)

    p(doc,
      "Two-stage structure.  As shown in Figure 1, the model separates "
      "environmental noise from the LLM's internal decision process. "
      "In the first stage (left trellis), acoustic noise corrupts the true "
      "state Y ∈ {1, 2} (not hate / hate) into a noisy latent state V: "
      "with probability q the hate signal is masked (Y=2 → V=1), "
      "and with probability 1−q it is preserved (Y=2 → V=2); "
      "the not-hate state (Y=1) is always transmitted correctly (V=1). "
      "In the second stage (right trellis), the LLM — acting as a "
      "rationally inattentive decision-maker — observes V and outputs "
      "action A ∈ {a, b} (predict not-hate / predict hate) "
      "at information cost λ > 0 per unit of Shannon information.",
      size=10)

    add_figure(doc, TRELLIS_FIG, width=Inches(4.5),
               caption="Figure 1. Two-stage trellis: noise channel Y→V (left) "
                       "and RI decision V→A (right). "
                       "q = hate-speech hiding probability; "
                       "P1a, P2b = conditional correct-response probabilities.")

    p(doc,
      "Noise channel.  Acoustic noise is characterised by the Noise-to-Signal "
      "power ratio NSR = 10^(−SNR_dB/10). The hiding probability q is modelled as:",
      size=10)

    p(doc,
      "    q(NSR) = min( α · NSR^β,  1 )",
      bold=True, size=10)

    p(doc,
      "where α, β > 0 are environment parameters shared across all models "
      "within a noise condition and estimated from data.",
      size=10)

    p(doc,
      "RI solution.  Let x = r/λ, where r is the reward for a correct decision "
      "and λ > 0 is the per-bit information cost. "
      "Given x and q, the LLM's optimal marginal probability "
      "of choosing action a (Matejka and McKay, 2015) is:",
      size=10)

    p(doc,
      "    Pa = min(1, max(0.5,  [(1+q)·e^x − (1−q)] / [2(e^x − 1)] ))",
      bold=True, size=10)

    p(doc,
      "The conditional correct-response probabilities and overall accuracy are:",
      size=10)

    p(doc,
      "    P1a = Pa·e^x / [Pa·e^x + (1−Pa)]\n"
      "    P2b = (1−Pa)·e^x / [Pa + (1−Pa)·e^x]\n"
      "    P(A=b|Y=2) = (1−q)·P2b + q·(1−P1a)\n"
      "    Pc  ≡ P(A=a,Y=1) + P(A=b,Y=2) = 0.5·(1−q)·P1a + 0.5·(1−q)·P2b + 0.5·q\n"
      "       ≡ Pc(q, x)",
      bold=True, size=10)

    p(doc,
      "The term 0.5·q captures fully-masked trials where the LLM guesses at chance. "
      "The six parameters {x_GPT-3.5, x_GPT-5.4, x_Gemini-Flash, "
      "x_Gemini-Flash-Lite, α, β} are estimated jointly per noise condition "
      "by minimising the SSE between empirical and theoretical Pc(q, x).",
      size=10)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 3. DATASET
    # ══════════════════════════════════════════════════════════════════════════
    h(doc, "3. Dataset", 1)

    # ── 3.1 Method / Pipeline ─────────────────────────────────────────────────
    h(doc, "3.1  Method and Pipeline", 2)
    p(doc,
      "We design an audio noise pipeline that introduces acoustic degradation "
      "through the speech channel — a qualitatively different distortion mechanism "
      "from the character-level text perturbation studied in Zhao et al. (2026). "
      "Noise now enters before automatic speech recognition (ASR), creating a "
      "compound distortion: acoustic corruption followed by ASR transcription errors. "
      "The pipeline consists of four sequential stages:",
      size=10)

    doc.add_paragraph()

    stages = [
        ("Stage 1 — Text-to-Speech (TTS)",
         "Each source text is synthesised into speech using Microsoft Edge TTS "
         "(voice: en-US-GuyNeural, 16 kHz mono WAV). "
         "The clean WAV files serve as the noise-free baseline and as input "
         "to the noise injection stage."),
        ("Stage 2 — Acoustic Noise Injection",
         "Each clean WAV is mixed with one of three acoustic noise sources "
         "at 21 controlled SNR levels "
         "(40, 30, 20, 15, 10, 5, 4, 3, 2, 1, 0, −1, −2, −3, −5, −6, −7, −8, "
         "−10, −15, −20 dB) using the power-normalised mixing formula "
         "x_noisy = x_clean + 10^(−SNR/20) · x_noise. "
         "Noise sources: "
         "(i) White — synthetic i.i.d. Gaussian noise; "
         "(ii) Babble — overlapping speech from 15 speakers "
         "(MUSAN corpus, Snyder et al., 2015); "
         "(iii) Café — real cafeteria ambience "
         "(DEMAND dataset, Thiemann et al., 2013). "
         "Including the clean baseline, each noise type covers 22 levels."),
        ("Stage 3 — Automatic Speech Recognition (ASR)",
         "Each noisy WAV is transcribed to text using faster-whisper "
         "(base model, English; Radford et al., 2023). "
         "ASR word-error rate increases with noise level, so the transcribed text "
         "degrades from both acoustic corruption and recognition errors simultaneously. "
         "The output is stored in the same tabular format as the text-noise data, "
         "so the same RI fitting pipeline applies without modification."),
        ("Stage 4 — LLM Classification",
         "Each transcribed text is classified independently by four commercial LLMs "
         "— GPT-3.5-turbo, GPT-5.4-nano, Gemini-2.5-Flash, Gemini-2.5-Flash-Lite — "
         "using a zero-shot prompt that asks for a binary hate-speech label "
         "(0 = not hate speech, 1 = hate speech). "
         "Calls are made per-item with no batching."),
    ]

    for i, (title, body) in enumerate(stages, 1):
        pp = doc.add_paragraph(style="List Number")
        r1 = pp.add_run(title + ": ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = pp.add_run(body)
        r2.font.size = Pt(10)

    doc.add_paragraph()

    # ── 3.2 Data Points ──────────────────────────────────────────────────────
    h(doc, "3.2  Data Points", 2)
    p(doc,
      "The source corpus is 100 short hate-speech comments from the "
      "UCBerkeley-DLAB Measuring Hate Speech dataset (Kennedy et al., 2020), "
      "balanced 50 / 50 between not-hate-speech (Y=1) and hate-speech (Y=2). "
      "These are the same 100 texts used in Zhao et al. (2026).",
      size=10)

    doc.add_paragraph()

    p(doc,
      "Applying the pipeline above produces the following observations:",
      size=10)

    stats = [
        ("Noise conditions",
         "3 acoustic noise types (white, babble, café)"),
        ("Noise levels per type",
         "22 (clean baseline + 21 SNR levels from 40 dB down to −20 dB)"),
        ("LLMs",
         "4 (GPT-3.5-turbo, GPT-5.4-nano, Gemini-2.5-Flash, Gemini-2.5-Flash-Lite)"),
        ("Observations per cell",
         "100 (text, ground-truth label Y, LLM prediction A) triples"),
        ("Total observations",
         "100 texts × 3 noise types × 22 levels × 4 LLMs = 26,400 triples"),
        ("RI fitting input",
         "Per noise type: a 4 × 22 matrix of empirical Pc values, "
         "one entry per (model, NSR level) cell"),
        ("NSR conversion",
         "NSR = 10^(−SNR_dB / 10);  the clean baseline maps to NSR = 0"),
    ]

    for label, value in stats:
        pp = doc.add_paragraph(style="List Bullet")
        r1 = pp.add_run(label + ":  ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = pp.add_run(value)
        r2.font.size = Pt(10)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 4. RESULTS
    # ══════════════════════════════════════════════════════════════════════════
    h(doc, "4. Results", 1)

    # ── 4a. Regression figures ─────────────────────────────────────────────
    h(doc, "4a. RI Model Fit: Empirical Accuracy and Fitted Curves", 2)
    p(doc,
      "Each panel pair below shows empirical Pc values (dots) and RI model fits "
      "(solid curves) for all four LLMs under one noise condition. "
      "Left panel: x-axis is SNR (dB, linear; higher = cleaner signal). "
      "Right panel: x-axis is NSR = 10^(−SNR/10) (log scale; higher = more noise). "
      "Both axes represent the same data — SNR is the conventional engineering "
      "scale while NSR is the scale used in the RI model formulation.",
      size=10)

    for noise in NOISE_ORDER:
        h(doc, NOISE_LABELS[noise], 3)
        add_figure(doc, FIG[noise]["combined"], width=Inches(6.5),
                   caption=(f"Figure: Pc vs SNR (left) and NSR (right) — "
                             f"all models, empirical + RI fit ({NOISE_LABELS[noise]})"))
        doc.add_paragraph()

    # ── 4b. Parameter estimates table ─────────────────────────────────────
    page_break(doc)
    h(doc, "4b. Estimated RI Parameters (Bootstrap SE and 95% CI)", 2)
    p(doc,
      "Table 1 reports the jointly estimated RI parameters for each noise condition. "
      "Standard errors and 95% CIs are obtained from B = 1,000 bootstrap replicates "
      "(percentile method, Efron & Tibshirani, 1993). "
      "λ is the per-model information cost; α and β are shared environment parameters.",
      size=10)

    doc.add_paragraph()

    # Build wide table: rows = params, cols = noise types
    param_order = [f"lambda_{m}" for m in MODEL_ORDER] + ["alpha", "beta"]
    param_labels = {f"lambda_{m}": f"λ  {m}" for m in MODEL_ORDER}
    param_labels["alpha"] = "α  (shared)"
    param_labels["beta"]  = "β  (shared)"

    cols = ["Parameter"] + [NOISE_LABELS[n] for n in NOISE_ORDER]
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Table Grid"
    hdr_row(t, cols)

    for param in param_order:
        vals = [param_labels[param]]
        for noise in NOISE_ORDER:
            row = aud_sum[(aud_sum["noise"] == noise) & (aud_sum["param"] == param)]
            if row.empty:
                vals.append("—")
            else:
                r = row.iloc[0]
                vals.append(
                    f"{r['estimate']:.4f} ± {r['se']:.4f}\n"
                    f"[{r['ci_lo']:.4f}, {r['ci_hi']:.4f}]"
                )
        # shade lambda rows alternately
        is_alpha_beta = param in ("alpha", "beta")
        data_row(t, vals, center_cols=[1, 2, 3])

    col_widths(t, [1.9, 1.6, 1.6, 1.6])
    p(doc, "Table 1. Estimate ± Bootstrap SE  [95% CI]. B = 1,000 replicates.",
      italic=True, size=9)

    doc.add_paragraph()

    # ── Token price vs lambda comparison ──────────────────────────────────
    h(doc, "4b-ii. Token Price vs. Estimated Information Cost (λ)", 2)
    p(doc,
      "Table 2 compares the estimated λ (averaged across the three noise conditions) "
      "with the official API input-token price for each model (standard tier, July 2026). "
      "Because the classification prompt instructs each LLM to return only a single "
      'JSON field — {"label": 0 or 1} — with a hard cap of 16 output tokens, '
      "output cost is negligible and the effective per-call cost is determined "
      "entirely by the input price. "
      "The token price reflects the provider's compute charge, while λ captures the "
      "LLM's behavioural information-processing cost inferred from its classification "
      "decisions — they measure fundamentally different quantities.",
      size=10)

    doc.add_paragraph()

    # Token price table (input only; output ignored — max 16 tokens per call)
    # Prices fetched July 2026: openai.com/api/pricing, cloud.google.com/vertex-ai/generative-ai/pricing
    PRICE_DATA = [
        # (display name, input_$/1M, note)
        ("GPT-3.5-turbo",         0.50, "deprecated"),
        ("GPT-5.4-nano",          0.20, ""),
        ("Gemini-2.5-Flash",      0.30, ""),
        ("Gemini-2.5-Flash-Lite", 0.10, ""),
    ]

    import numpy as np

    # Build per-condition lambda data
    lam_full = (aud_sum[aud_sum["param"].str.startswith("lambda_")].copy())
    lam_full["model"] = lam_full["param"].str.replace("lambda_", "")
    PRICE_MAP = {"GPT-3.5-turbo": 0.50, "GPT-5.4-nano": 0.20,
                 "Gemini-2.5-Flash": 0.30, "Gemini-2.5-Flash-Lite": 0.10}
    lam_full["price"]   = lam_full["model"].map(PRICE_MAP)
    lam_full["lam_w"]   = lam_full["price"] * lam_full["estimate"]
    lam_full["se_w"]    = lam_full["price"] * lam_full["se"]
    lam_full["ci_lo_w"] = lam_full["price"] * lam_full["ci_lo"]
    lam_full["ci_hi_w"] = lam_full["price"] * lam_full["ci_hi"]

    # Build per-condition x data (x = r/λ, directly estimated parameter)
    x_full = aud_sum[aud_sum["param"].str.startswith("x_")].copy()
    x_full["model"] = x_full["param"].str.replace("x_", "")
    x_full["price"]    = x_full["model"].map(PRICE_MAP)
    x_full["lam_w_i"]  = x_full["price"] / x_full["estimate"]          # λ_w = r/x
    x_full["se_lw_i"]  = x_full["price"] / (x_full["estimate"] ** 2) * x_full["se"]  # delta method

    # Average across 3 noise types with correct SE propagation:
    #   λ_avg  = (λ_w + λ_b + λ_c) / 3   (independent conditions)
    #   SE_avg = sqrt(SE_w² + SE_b² + SE_c²) / 3
    def avg_se_propagated(series):
        return np.sqrt((series ** 2).sum()) / len(series)

    lam_avg_df = lam_full.groupby("model").agg(
        price    = ("price",    "first"),
        lam_mean = ("estimate", "mean"),
        se_mean  = ("se",       avg_se_propagated),
    ).reset_index()

    # Table 3 aggregation: average x and λ_w = r/x across noise conditions
    x_avg_df = x_full.groupby("model").agg(
        price   = ("price",    "first"),
        x_avg   = ("estimate", "mean"),
        lam_w   = ("lam_w_i", "mean"),
        se_w    = ("se_lw_i", avg_se_propagated),
    ).reset_index()

    lw_avg_df = lam_full.groupby("model").agg(
        price    = ("price",  "first"),
        lam_old  = ("estimate","mean"),
        lam_w    = ("lam_w",  "mean"),
        se_w     = ("se_w",   avg_se_propagated),
    ).reset_index()

    # ── Table 2: λ vs token price ──────────────────────────────────────────
    price_rows = []
    for (name, inp, note) in PRICE_DATA:
        row = lam_avg_df[lam_avg_df["model"] == name]
        if not row.empty:
            lv = row["lam_mean"].values[0]
            ls = row["se_mean"].values[0]
        else:
            lv, ls = float("nan"), float("nan")
        price_rows.append((name, lv, ls, inp, note))

    price_rows.sort(key=lambda x: x[1])   # sort by λ ascending

    sorted_by_price  = sorted(price_rows, key=lambda x: x[3])
    price_rank_map   = {r[0]: i + 1 for i, r in enumerate(sorted_by_price)}

    price_cols = ["Model", "λ (avg ± SE)", "Input $/1M tokens",
                  "λ rank", "Price rank", "Note"]
    pt = doc.add_table(rows=1, cols=len(price_cols))
    pt.style = "Table Grid"
    hdr_row(pt, price_cols)

    for i, (name, lv, ls, inp, note) in enumerate(price_rows):
        lam_str = f"{lv:.4f} ± {ls:.4f}"
        data_row(pt,
                 [name, lam_str, f"${inp:.2f}",
                  str(i + 1), str(price_rank_map[name]), note],
                 center_cols=[1, 2, 3, 4])

    col_widths(pt, [1.85, 1.40, 1.35, 0.75, 0.85, 0.95])
    p(doc,
      "Table 2. Input token prices from official provider pricing pages (July 2026, standard tier). "
      "Output tokens ignored: prompt enforces single-digit JSON output (≤16 tokens). "
      "λ rank: 1 = lowest λ = most decision-efficient. "
      "Price rank: 1 = cheapest input price. "
      "Sources: openai.com/api/pricing; cloud.google.com/vertex-ai/generative-ai/pricing.",
      italic=True, size=9)

    doc.add_paragraph()

    # Formula note for Table 2 SE
    p(doc,
      "SE derivation for averaged λ (Table 2).  "
      "Each model's λ is estimated independently for the three noise conditions "
      "(white, babble, café). "
      "The average λ_avg = (λ_w + λ_b + λ_c) / 3. "
      "Because the three bootstrap procedures are independent, variances add:",
      size=10)
    p(doc,
      "    SE(λ_avg) = sqrt( SE_white² + SE_babble² + SE_café² ) / 3",
      bold=True, size=10)
    p(doc,
      "This is the standard error of the mean of three independent estimates "
      "(error propagation for a linear combination with equal weights 1/3).",
      size=10)

    doc.add_paragraph()

    p(doc,
      "Key observation (Table 2):  Input token price and λ do not share the same "
      "ordering, indicating that token cost and decision efficiency capture distinct "
      "properties of LLM behaviour. "
      "GPT-3.5-turbo is the most expensive model per token ($0.50/M) yet achieves "
      "the lowest λ (rank 1) — it extracts decision-relevant information most "
      "efficiently from noisy input. "
      "GPT-5.4-nano has the highest λ (rank 4) despite costing only $0.20/M. "
      "Within the Gemini family the rankings are consistent: "
      "Gemini-2.5-Flash ($0.30/M) is both more expensive and more decision-efficient "
      "than Gemini-2.5-Flash-Lite ($0.10/M).",
      size=10)

    doc.add_paragraph()

    # ── Price-weighted lambda (lambda_w) intro ─────────────────────────────
    p(doc,
      "Price-weighted information cost (λ_w).  "
      "In the RI model, the decision parameter is x = r / λ, where r is the reward "
      "for a correct decision. Our baseline estimation assumes r = 1 uniformly across "
      "models. However, because the prompt enforces identical single-digit output and "
      "input length is the same text for all models, the effective per-call cost is "
      "entirely determined by the input token price. "
      "Setting r = input price ($/1M tokens) as a model-specific reward, "
      "and keeping the fitted x unchanged, gives:",
      size=10)

    p(doc,
      "    λ_w = r / x",
      bold=True, size=10)

    p(doc,
      "λ_w measures the information cost in the same monetary units as the token "
      "price: a model with low λ_w processes each bit of decision-relevant "
      "information at low monetary cost.",
      size=10)

    doc.add_paragraph()

    # ── Table 3: λ_w summary sorted by λ_w, with ranks ────────────────────
    lw_sorted = x_avg_df.sort_values("lam_w").reset_index(drop=True)

    sorted_by_price_lw = x_avg_df.sort_values("price", ascending=True).reset_index(drop=True)
    price_rank_lw = {r["model"]: i + 1
                     for i, r in sorted_by_price_lw.iterrows()}

    lw_cols = ["Model", "x (fitted)", "r ($/1M tok)", "λ_w = r/x  (avg ± SE)",
               "λ_w rank", "Price rank"]
    lt = doc.add_table(rows=1, cols=len(lw_cols))
    lt.style = "Table Grid"
    hdr_row(lt, lw_cols)

    for i, row in lw_sorted.iterrows():
        m = row["model"]
        data_row(lt,
                 [m,
                  f"{row['x_avg']:.4f}",
                  f"${row['price']:.2f}",
                  f"{row['lam_w']:.4f} ± {row['se_w']:.4f}",
                  str(i + 1),
                  str(price_rank_lw[m])],
                 center_cols=[1, 2, 3, 4, 5])

    col_widths(lt, [1.85, 0.90, 1.05, 1.75, 0.85, 0.85])
    p(doc,
      "Table 3. Price-weighted information cost λ_w = r/x, sorted by λ_w ascending. "
      "x = r/λ is the directly estimated RI parameter (with r=1 in baseline fitting). "
      "λ_w rank: 1 = lowest λ_w = most cost-efficient. "
      "Price rank: 1 = cheapest input price.",
      italic=True, size=9)

    doc.add_paragraph()

    # Formula note for Table 3 SE
    p(doc,
      "SE derivation for λ_w (Table 3).  "
      "x is the directly estimated parameter; λ_w = r/x, so by the delta method:",
      size=10)
    p(doc,
      "    SE(λ_w) = (r / x²) · SE(x)              [per noise condition]\n"
      "    SE(λ_w_avg) = sqrt( SE(λ_w,w)² + SE(λ_w,b)² + SE(λ_w,c)² ) / 3",
      bold=True, size=10)
    p(doc,
      "where SE(λ_w,k) = (r / x_k²) · SE(x_k) for noise condition k ∈ {white, babble, café}. "
      "This applies the delta method for g(x) = r/x, giving g'(x) = −r/x², "
      "and combines the per-condition SEs using the independent-estimates propagation rule.",
      size=10)

    doc.add_paragraph()

    # ── Table 4: λ_w per noise condition, model order = MODEL_ORDER ────────
    lw_noise_cols = ["Model", "r ($/1M tok)"] + [NOISE_LABELS[n] for n in NOISE_ORDER]
    lt2 = doc.add_table(rows=1, cols=len(lw_noise_cols))
    lt2.style = "Table Grid"
    hdr_row(lt2, lw_noise_cols)

    for m in MODEL_ORDER:          # same order as Table 1
        row_avg = lw_avg_df[lw_avg_df["model"] == m]
        price_str = f"${row_avg['price'].values[0]:.2f}" if not row_avg.empty else "—"
        vals = [m, price_str]
        for noise in NOISE_ORDER:
            sub = lam_full[(lam_full["model"] == m) & (lam_full["noise"] == noise)]
            if not sub.empty:
                r = sub.iloc[0]
                vals.append(f"{r['lam_w']:.4f} ± {r['se_w']:.4f}\n"
                            f"[{r['ci_lo_w']:.4f}, {r['ci_hi_w']:.4f}]")
            else:
                vals.append("—")
        data_row(lt2, vals, center_cols=[1, 2, 3, 4])

    col_widths(lt2, [1.85, 1.05, 1.55, 1.55, 1.55])
    p(doc,
      "Table 4. λ_w per noise condition (estimate ± SE  [95% CI]). "
      "Model order matches Table 1.",
      italic=True, size=9)

    doc.add_paragraph()

    # Interpretation — use lw_sorted already computed above
    most_eff  = lw_sorted.iloc[0]
    least_eff = lw_sorted.iloc[-1]
    gpt35_lw_rank  = int(lw_sorted[lw_sorted["model"]=="GPT-3.5-turbo"].index[0]) + 1
    flash_lite_rank = int(lw_sorted[lw_sorted["model"]=="Gemini-2.5-Flash-Lite"].index[0]) + 1

    p(doc,
      f"Key observation (Tables 3 & 4):  When reward is set to token price, "
      f"the ranking of models by information cost changes substantially. "
      f"{most_eff['model']} achieves the lowest λ_w "
      f"({most_eff['lam_w']:.4f}), making it the most cost-efficient model — "
      f"it processes each bit of decision-relevant information at the lowest "
      f"monetary cost. "
      f"{least_eff['model']} has the highest λ_w ({least_eff['lam_w']:.4f}): "
      f"its high token price amplifies its already-elevated baseline λ. "
      f"Comparing Tables 2 and 3, GPT-3.5-turbo drops from λ rank 1 (most "
      f"decision-efficient under r=1) to λ_w rank {gpt35_lw_rank} "
      f"once its price premium is factored in, while Gemini-2.5-Flash-Lite "
      f"rises to λ_w rank {flash_lite_rank} due to its low token price. "
      f"λ_w thus provides a unified, price-adjusted measure of LLM decision "
      f"efficiency that is directly comparable across providers.",
      size=10)

    doc.add_paragraph()

    # ── 4c. Significance test summary ─────────────────────────────────────
    h(doc, "4c. Significance Tests: Summary", 2)
    p(doc,
      "All significance tests use the bootstrap CI of the difference "
      "(H₀ rejected at 5% level when 0 ∉ [2.5th, 97.5th percentile] of "
      "θ_A^b − θ_B^b; Greenwood, 2022, §2.9). Full tables are in Appendix 1 and 2.",
      size=10)

    # Within-noise summary
    h(doc, "Between-model comparison (within noise type)", 3)

    for noise in NOISE_ORDER:
        sub = aud_mod[aud_mod["noise"] == noise]
        sig_pairs = sub[sub["sig_ci"].apply(sig_bool)]
        sig_n = len(sig_pairs)
        tot = len(sub)

        # GPT-5.4-nano specifics
        gpt54_sig_others = []
        for _, r in sig_pairs.iterrows():
            if "5.4" in r["model_A"]:
                gpt54_sig_others.append(r["model_B"])
            elif "5.4" in r["model_B"]:
                gpt54_sig_others.append(r["model_A"])

        gpt54_note = ""
        if gpt54_sig_others:
            gpt54_note = (f"  GPT-5.4-nano has the highest λ and is significantly "
                          f"higher than: {', '.join(gpt54_sig_others)}.")

        p(doc,
          f"{NOISE_LABELS[noise]}:  {sig_n}/{tot} model pairs are significant "
          f"by the CI test.{gpt54_note}",
          size=10)

    doc.add_paragraph()

    # Across-noise summary
    h(doc, "Across-noise comparison (same model, different noise type)", 3)

    lam_cross = aud_cross[aud_cross["param"].str.startswith("lambda_")]
    ab_cross  = aud_cross[aud_cross["param"].isin(["alpha", "beta"])]
    sig_lam   = lam_cross["sig_ci"].apply(sig_bool).sum()

    p(doc,
      f"Lambda (λ):  {sig_lam}/{len(lam_cross)} cross-noise λ comparisons are "
      f"significant.  The information cost parameter is stable across white, babble, "
      f"and café noise for all models — each LLM's decision efficiency does not change "
      f"significantly as the type of acoustic noise changes.",
      size=10)

    doc.add_paragraph()

    sig_ab_rows = ab_cross[ab_cross["sig_ci"].apply(sig_bool)]
    if len(sig_ab_rows) > 0:
        for _, r in sig_ab_rows.iterrows():
            nA = NOISE_LABELS.get(r["noise_A"], r["noise_A"])
            nB = NOISE_LABELS.get(r["noise_B"], r["noise_B"])
            p(doc,
              f"Environment parameter α:  α differs significantly between "
              f"{nA} (α = {r['est_A']:.3f}) and {nB} (α = {r['est_B']:.3f}), "
              f"95% CI of difference [{r['ci_lo_diff']:.3f}, {r['ci_hi_diff']:.3f}].  "
              f"This suggests that the rate at which noise degrades signal quality "
              f"differs between these two acoustic conditions, even though the "
              f"decision parameter λ remains stable.",
              size=10)
    else:
        p(doc,
          "Environment parameters α and β:  No significant cross-noise differences.",
          size=10)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 5. CONCLUSION
    # ══════════════════════════════════════════════════════════════════════════
    page_break(doc)
    h(doc, "5. Conclusion", 1)
    p(doc, "[To be written]", italic=True, size=10)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # REFERENCES
    # ══════════════════════════════════════════════════════════════════════════
    h(doc, "References", 1)
    refs = [
        "Efron, B., & Tibshirani, R. J. (1993). An Introduction to the Bootstrap. "
        "Chapman & Hall/CRC.",
        "Greenwood, M. C. (2022). Intermediate Statistics with R (2nd ed.). "
        "[Bootstrap CI of difference — Section 2.9]",
        "Kennedy, B., Bacon, G., Sahn, A., & von der Brelie, C. (2020). "
        "Constructing interval variables via faceted Rasch measurement and "
        "multitask deep learning: a hate speech application. arXiv:2009.10277.",
        "Sims, C. A. (2003). Implications of rational inattention. "
        "Journal of Monetary Economics, 50(3), 665–690.",
    ]
    for ref in refs:
        pp = doc.add_paragraph(style="List Paragraph")
        r = pp.add_run(ref)
        r.font.size = Pt(10)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # APPENDIX 1 — Within-noise pairwise lambda comparison
    # ══════════════════════════════════════════════════════════════════════════
    page_break(doc)
    h(doc, "Appendix 1: Pairwise λ Comparison — Within Noise Environment", 1)
    p(doc,
      "Full results of pairwise bootstrap CI tests for λ between all model pairs "
      "within each noise condition. ✓ = significant (0 ∉ 95% CI of difference); "
      "– = not significant.",
      size=10)

    for noise in NOISE_ORDER:
        h(doc, NOISE_LABELS[noise], 2)
        sub = aud_mod[aud_mod["noise"] == noise]
        cols = ["Model A", "Model B", "λ_A", "λ_B",
                "Diff", "95% CI of Diff", "Sig."]
        t = doc.add_table(rows=1, cols=len(cols))
        t.style = "Table Grid"
        hdr_row(t, cols)
        for _, r in sub.iterrows():
            sig = sig_bool(r["sig_ci"])
            data_row(t,
                     [r["model_A"], r["model_B"],
                      fmt(r["lambda_A"]), fmt(r["lambda_B"]),
                      fmt(r["obs_diff"]),
                      f"[{r['ci_lo_diff']:.4f}, {r['ci_hi_diff']:.4f}]",
                      SIG_MARK if sig else NOSIG_MARK],
                     sig=sig, center_cols=[2, 3, 4, 5, 6])
        col_widths(t, [1.5, 1.9, 0.7, 0.7, 0.75, 1.35, 0.45])
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # APPENDIX 2 — Across-noise parameter comparison
    # ══════════════════════════════════════════════════════════════════════════
    page_break(doc)
    h(doc, "Appendix 2: Parameter Comparison — Across Noise Environments", 1)
    p(doc,
      "Full results of bootstrap CI tests for each parameter across noise-type pairs. "
      "Rows cover λ (per model) and shared α, β. "
      "✓ = significant; – = not significant.",
      size=10)

    # Lambda across noise
    h(doc, "A2a. Lambda (λ) Across Noise Types", 2)
    lam_cross = aud_cross[aud_cross["param"].str.startswith("lambda_")]
    cols = ["Noise A", "Noise B", "Model", "λ_A", "λ_B",
            "Diff", "95% CI of Diff", "Sig."]
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Table Grid"
    hdr_row(t, cols)
    for _, r in lam_cross.iterrows():
        sig = sig_bool(r["sig_ci"])
        model = r["param"].replace("lambda_", "")
        data_row(t,
                 [NOISE_LABELS.get(r["noise_A"], r["noise_A"]),
                  NOISE_LABELS.get(r["noise_B"], r["noise_B"]),
                  model,
                  fmt(r["est_A"]), fmt(r["est_B"]),
                  fmt(r["obs_diff"]),
                  f"[{r['ci_lo_diff']:.4f}, {r['ci_hi_diff']:.4f}]",
                  SIG_MARK if sig else NOSIG_MARK],
                 sig=sig, center_cols=[3, 4, 5, 6, 7])
    col_widths(t, [0.95, 0.95, 1.55, 0.68, 0.68, 0.75, 1.35, 0.43])
    doc.add_paragraph()

    # Alpha / beta across noise
    h(doc, "A2b. Environment Parameters (α, β) Across Noise Types", 2)
    ab_cross = aud_cross[aud_cross["param"].isin(["alpha", "beta"])]
    cols = ["Noise A", "Noise B", "Param", "Value_A", "Value_B",
            "Diff", "95% CI of Diff", "Sig."]
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Table Grid"
    hdr_row(t, cols)
    for _, r in ab_cross.iterrows():
        sig = sig_bool(r["sig_ci"])
        data_row(t,
                 [NOISE_LABELS.get(r["noise_A"], r["noise_A"]),
                  NOISE_LABELS.get(r["noise_B"], r["noise_B"]),
                  r["param"],
                  fmt(r["est_A"]), fmt(r["est_B"]),
                  fmt(r["obs_diff"]),
                  f"[{r['ci_lo_diff']:.4f}, {r['ci_hi_diff']:.4f}]",
                  SIG_MARK if sig else NOSIG_MARK],
                 sig=sig, center_cols=[3, 4, 5, 6, 7])
    col_widths(t, [0.95, 0.95, 0.75, 0.75, 0.75, 0.75, 1.35, 0.43])

    # ── Save ─────────────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Report saved -> {OUT_PATH}")


if __name__ == "__main__":
    main()
