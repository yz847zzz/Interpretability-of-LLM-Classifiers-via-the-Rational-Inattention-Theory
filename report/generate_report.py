"""
Generate a Word report summarising the audio noise RI experiment.

Usage:
    python generate_report.py
Output:
    Dataset/results/audio_noise_report.docx
"""

import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

RESULTS_DIR = "./Dataset/results"
NOISE_TYPES = ["babble", "cafe", "white"]
OUT_PATH    = os.path.join(RESULTS_DIR, "audio_noise_report.docx")

# ── helpers ──────────────────────────────────────────────────────────────────

def heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def add_image(doc, path, width=Inches(5.5), caption=None):
    if not os.path.exists(path):
        para(doc, f"[Image not found: {path}]")
        return
    doc.add_picture(path, width=width)
    if caption:
        c = doc.add_paragraph(caption)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.runs[0].font.size = Pt(9)
        c.runs[0].italic = True

def df_to_table(doc, df, style="Table Grid"):
    has_index = bool(df.index.name)   # only show index column if it has a name
    t = doc.add_table(rows=1 + len(df), cols=len(df.columns) + (1 if has_index else 0))
    t.style = style
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header
    hdr = t.rows[0].cells
    col_offset = 0
    if has_index:
        hdr[0].text = str(df.index.name or "")
        col_offset = 1
    for j, c in enumerate(df.columns):
        hdr[j + col_offset].text = str(c)
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # rows
    for i, (idx, row) in enumerate(df.iterrows()):
        cells = t.rows[i + 1].cells
        if has_index:
            cells[0].text = str(idx)
        for j, val in enumerate(row):
            cells[j + col_offset].text = str(val)

    return t

def add_images_2col(doc, paths, width=Inches(3.0)):
    """Insert images two per row using a borderless 2-column table."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pairs = [paths[i:i+2] for i in range(0, len(paths), 2)]
    for pair in pairs:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        # remove all borders
        for row in tbl.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement("w:tcBorders")
                for side in ("top","left","bottom","right","insideH","insideV"):
                    border = OxmlElement(f"w:{side}")
                    border.set(qn("w:val"), "none")
                    tcBorders.append(border)
                tcPr.append(tcBorders)

        for col_idx, img_path in enumerate(pair):
            cell = tbl.cell(0, col_idx)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if os.path.exists(img_path):
                run = cell.paragraphs[0].add_run()
                run.add_picture(img_path, width=width)
            else:
                cell.paragraphs[0].add_run(f"[Not found: {os.path.basename(img_path)}]")
        doc.add_paragraph()

# ── document ─────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # title
    title = doc.add_heading("Audio Noise RI Experiment — Results Summary", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Section 1: Data Generation ────────────────────────────────────────────
    heading(doc, "1. Data Generation Pipeline", 1)

    para(doc, (
        "The experiment uses 100 hate-speech texts with binary labels (50 no-hate / 50 hate) "
        "sampled from the UCBerkeley-DLAB Measuring Hate Speech dataset "
        "(Kennedy et al., 2020). The full audio noise pipeline consists of four stages:"
    ))

    steps = [
        ("Stage 1 — Text → Audio (TTS)",
         "Each text is synthesised into speech using edge-tts (voice: en-US-GuyNeural) "
         "and saved as a 16 kHz mono WAV file (100 files total)."),
        ("Stage 2 — Audio + Noise Injection",
         "Each clean WAV is mixed with one of three real-world noise recordings at 21 SNR levels "
         "(40, 30, 20, 15, 10, 5, 4, 3, 2, 1, 0, −1, −2, −3, −5, −6, −7, −8, −10, −15, −20 dB) "
         "using the power-normalised mixing formula. "
         "Noise types: white (synthetic Gaussian), babble (MUSAN speech overlay, 15 speakers), "
         "cafe (DEMAND CAFE real cafeteria recording). "
         "A noise-free baseline (NSR = 0) is taken directly from the original text."),
        ("Stage 3 — Audio → Text (ASR)",
         "Noisy WAVs are transcribed using faster-whisper (base model) to produce "
         "one noisy-text CSV per (noise type × SNR level), matching the format of "
         "the text-noise pipeline."),
        ("Stage 4 — LLM Classification",
         "Each noisy text is classified independently by four LLMs: "
         "GPT-3.5-turbo, GPT-5.4-nano, Gemini-2.5-Flash, Gemini-2.5-Flash-Lite. "
         "Calls are made per-item (no batching) to avoid cross-item influence. "
         "Prompt instructs binary hate-speech classification (0 = no hate, 1 = hate), "
         "tie-break: choose 0 if unsure."),
    ]
    for title_s, body in steps:
        p = doc.add_paragraph(style="List Number")
        p.add_run(title_s + ": ").bold = True
        p.add_run(body)

    para(doc, (
        "\nThe NSR (Noise-to-Signal power ratio) is computed as NSR = N/S = 10^(−SNR_dB/10). "
        "The noise channel confusion probability is modelled as "
        "q(NSR) = min(α·NSR^β, 1), mirroring the text-noise formulation q(p′) = min(α·p′^β, 1)."
    ))

    # ── Section 2: Results ────────────────────────────────────────────────────
    heading(doc, "2. Results", 1)

    para(doc, (
        "Results are presented for three noise conditions. Each subsection shows: "
        "Pc vs SNR, Pc vs NSR, regression-only comparison, q vs SNR, q vs NSR "
        "and individual model fits."
    ))

    plot_groups = [
        ("SNR_Plot/all_models_fit.png",              "P_correct vs SNR — all models (empirical + fit)"),
        ("SNR_Plot/all_models_fit_regression_only.png", "P_correct vs SNR — regression lines only"),
        ("ri_fit/all_models_fit.png",                "P_correct vs NSR — all models (empirical + fit)"),
        ("ri_fit/all_models_fit_regression_only.png",   "P_correct vs NSR — regression lines only"),
        ("SNR_Plot/q_vs_SNR.png",                    "Hate speech hiding rate vs SNR"),
        ("ri_fit/q_vs_NSR.png",                      "Hate speech hiding rate vs NSR"),
    ]

    for noise in NOISE_TYPES:
        heading(doc, f"2.{NOISE_TYPES.index(noise)+1}  Noise type: {noise.capitalize()}", 2)
        noise_dir = os.path.join(RESULTS_DIR, f"audio_{noise}")

        for rel_path, caption in plot_groups:
            full_path = os.path.join(noise_dir, rel_path)
            add_image(doc, full_path, width=Inches(5.2),
                      caption=f"Figure: {caption} ({noise})")
            doc.add_paragraph()

        # individual model fits (Pc vs NSR) — 2 per row
        para(doc, "Individual model fits (Pc vs NSR):", bold=True)
        model_tags = [
            "GPT3.5turbo_fit.png",
            "GPT5.4nano_fit.png",
            "Gemini2.5Flash_fit.png",
            "Gemini2.5FlashLite_fit.png",
        ]
        model_paths = [os.path.join(noise_dir, "ri_fit", tag) for tag in model_tags]
        add_images_2col(doc, model_paths, width=Inches(3.0))

        doc.add_page_break()

    # ── Section 3: Parameter Estimation ──────────────────────────────────────
    heading(doc, "3. RI Model Parameter Estimation", 1)

    para(doc, (
        "Parameters are estimated by jointly minimising SSE across all four LLMs "
        "and all NSR levels within each noise condition. "
        "Point estimates are from the full-data fit. "
        "Standard errors and 95% confidence intervals are obtained via nonparametric "
        "bootstrap (B = 1000 resamples; Efron & Tibshirani, 1993)."
    ))

    summary_path = os.path.join(RESULTS_DIR, "bootstrap", "bootstrap_summary.csv")
    if os.path.exists(summary_path):
        df_s = pd.read_csv(summary_path)
        # filter lambda_ and alpha/beta, pivot to param × noise
        df_s = df_s[df_s["param"].str.startswith("lambda_") | df_s["param"].isin(["alpha","beta"])]

        rename = {
            "lambda_GPT-3.5-turbo":         "λ  GPT-3.5-turbo",
            "lambda_GPT-5.4-nano":           "λ  GPT-5.4-nano",
            "lambda_Gemini-2.5-Flash":       "λ  Gemini-2.5-Flash",
            "lambda_Gemini-2.5-Flash-Lite":  "λ  Gemini-2.5-Flash-Lite",
            "alpha":                          "α  (shared)",
            "beta":                           "β  (shared)",
        }
        row_order = list(rename.values())

        # build wide table: Parameter | babble_est | babble_SE | cafe_est | cafe_SE | ...
        cols = ["Parameter"]
        for n in [nn for nn in NOISE_TYPES if nn in df_s["noise"].unique()]:
            cols += [f"{n} est.", f"{n} SE"]

        rows_out = []
        for orig, label in rename.items():
            row = {"Parameter": label}
            for n in NOISE_TYPES:
                sub = df_s[(df_s["param"] == orig) & (df_s["noise"] == n)]
                if not sub.empty:
                    row[f"{n} est."] = f"{sub['estimate'].iloc[0]:.4f}"
                    row[f"{n} SE"]   = f"{sub['se'].iloc[0]:.4f}"
                else:
                    row[f"{n} est."] = ""
                    row[f"{n} SE"]   = ""
            rows_out.append(row)

        table_df = pd.DataFrame(rows_out, columns=[c for c in cols if c in rows_out[0]])
        table_df.index.name = None

        para(doc, "\nTable 1: RI parameters (estimate and SE from bootstrap, B=1000)", bold=True)
        df_to_table(doc, table_df)
    else:
        para(doc, "[bootstrap_summary.csv not found — run bootstrap_ri.py first]")

    doc.add_paragraph()

    # ── Section 4: Significance Tests ─────────────────────────────────────────
    heading(doc, "4. Pairwise Significance Tests", 1)

    para(doc, (
        "Two tests are reported for each pairwise comparison across noise conditions:\n"
        "(1) Z-test: z = (est_A − est_B) / √(SE_A² + SE_B²), two-sided p-value from "
        "standard normal (valid when bootstrap distribution is approximately normal). "
        "SE from bootstrap (B = 1000).\n"
        "(2) Bootstrap 95% CI for the difference (est_A − est_B): "
        "sig_CI = True when the CI excludes 0, i.e., the two noise conditions are "
        "significantly different at the 5% level regardless of distributional assumptions."
    ))

    pairs_path = os.path.join(RESULTS_DIR, "bootstrap", "pairwise_tests.csv")
    if os.path.exists(pairs_path):
        df_p = pd.read_csv(pairs_path)
        df_p = df_p[df_p["param"].str.startswith("lambda_") | df_p["param"].isin(["alpha","beta"])]
        df_p = df_p[["noise_A","noise_B","param","est_A","est_B","obs_diff",
                     "z_stat","p_z","sig_z_0.05",
                     "ci_lo_diff","ci_hi_diff","sig_ci"]]
        df_p.columns = ["Noise A","Noise B","Parameter","Est A","Est B","Diff",
                        "z","p(z)","sig(z)",
                        "CI lo","CI hi","sig(CI)"]
        df_p = df_p.reset_index(drop=True)

        para(doc, "\nTable 2: Pairwise significance tests (same parameter across noise types)", bold=True)
        t = df_to_table(doc, df_p)

        # highlight rows significant by either test
        for i, (_, row) in enumerate(df_p.iterrows()):
            if row["sig(z)"] or row["sig(CI)"]:
                for cell in t.rows[i + 1].cells:
                    cell.paragraphs[0].runs[0].font.bold = True if cell.paragraphs[0].runs else None
    else:
        para(doc, "[pairwise_tests.csv not found — run bootstrap_ri.py first]")

    doc.add_paragraph()

    # Table 3: pairwise model lambda comparisons within each noise type
    para(doc, (
        "Table 3 compares the λ (attention cost) values between all pairs of models "
        "within the same noise condition. The same two tests are used: z-test and bootstrap 95% CI "
        "for the difference (λ_A − λ_B). A significant result indicates that the two models "
        "allocate statistically different levels of attention under that noise type."
    ))

    model_pairs_path = os.path.join(RESULTS_DIR, "bootstrap", "pairwise_model_tests.csv")
    if os.path.exists(model_pairs_path):
        df_m = pd.read_csv(model_pairs_path)
        df_m = df_m[["noise","model_A","model_B","lambda_A","lambda_B","obs_diff",
                     "z_stat","p_z","sig_z_0.05",
                     "ci_lo_diff","ci_hi_diff","sig_ci"]]
        df_m.columns = ["Noise","Model A","Model B","λ_A","λ_B","Diff",
                        "z","p(z)","sig(z)",
                        "CI lo","CI hi","sig(CI)"]
        # sort by noise type order
        noise_order = {n: i for i, n in enumerate(NOISE_TYPES)}
        df_m["_ord"] = df_m["Noise"].map(noise_order)
        df_m = df_m.sort_values("_ord").drop(columns="_ord").reset_index(drop=True)

        para(doc, "\nTable 3: Pairwise λ comparisons between models within each noise type", bold=True)
        t3 = df_to_table(doc, df_m)

        # bold-highlight significant rows
        for i, (_, row) in enumerate(df_m.iterrows()):
            if row["sig(z)"] or row["sig(CI)"]:
                for cell in t3.rows[i + 1].cells:
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.bold = True
    else:
        para(doc, "[pairwise_model_tests.csv not found — run bootstrap_ri.py first]")

    # ── save ──────────────────────────────────────────────────────────────────
    doc.save(OUT_PATH)
    print(f"Report saved -> {OUT_PATH}")


if __name__ == "__main__":
    build()
