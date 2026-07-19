"""
Generate Word report summarising bootstrap CI results for both noise environments.

Analysis 1: Within each noise environment, pairwise lambda comparison across models.
Analysis 2: Across noise environments (audio only), same model lambda comparison.

Sources:
  data/results/text_noise_set2/bootstrap_summary.csv
  data/results/text_noise_set2/pairwise_lambda_tests.csv
  Dataset/results/bootstrap/bootstrap_summary.csv
  Dataset/results/bootstrap/pairwise_model_tests.csv   <- within-noise model pairs
  Dataset/results/bootstrap/pairwise_tests.csv         <- across-noise same-model

Output:
  Dataset/results/final_bootstrap_report.docx
"""

import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ─────────────────────────────────────────────────────────────────────
TEXT_SUMMARY   = "./data/results/text_noise_set2/bootstrap_summary.csv"
TEXT_PAIRWISE  = "./data/results/text_noise_set2/pairwise_lambda_tests.csv"
AUDIO_SUMMARY  = "./Dataset/results/bootstrap/bootstrap_summary.csv"
AUDIO_MODEL    = "./Dataset/results/bootstrap/pairwise_model_tests.csv"
AUDIO_CROSS    = "./Dataset/results/bootstrap/pairwise_tests.csv"
OUT_PATH       = "./Dataset/results/final_bootstrap_report_v3.docx"

NOISE_LABELS   = {"white": "White Noise", "babble": "Babble Noise", "cafe": "Café Noise"}
SIG_MARK       = "✓"
NOSIG_MARK     = "–"


# ── docx helpers ──────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_col_widths(table, widths_inches):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_inches):
                cell.width = Inches(widths_inches[i])


def add_header_row(table, cols, bg="2F5496"):
    row = table.rows[0]
    for i, text in enumerate(cols):
        cell = row.cells[i]
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, bg)


def add_data_row(table, values, sig=False, align_center=None):
    align_center = align_center or []
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        if i in align_center:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if sig and i == len(values) - 1:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1F, 0x7A, 0x1F)
    if sig:
        for cell in row.cells:
            set_cell_bg(cell, "EBF5EB")


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def fmt_ci(lo, hi):
    return f"[{lo:.3f}, {hi:.3f}]"


def sig_bool(val):
    if isinstance(val, bool):
        return val
    return str(val).strip().lower() == "true"


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    # Load data
    text_sum  = pd.read_csv(TEXT_SUMMARY)
    text_pair = pd.read_csv(TEXT_PAIRWISE)
    aud_sum   = pd.read_csv(AUDIO_SUMMARY)
    aud_mod   = pd.read_csv(AUDIO_MODEL)
    aud_cross = pd.read_csv(AUDIO_CROSS)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Title ──────────────────────────────────────────────────────────────────
    title = doc.add_heading(
        "Bootstrap Inference for Rational Inattention Parameters", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(doc,
         "This report presents 95% bootstrap confidence intervals (percentile method, "
         "B = 1000 replicates) for the Rational Inattention (RI) model parameters "
         "estimated across two noise environments: text noise and audio noise. "
         "Significance is assessed via the CI of difference: "
         "H₀: θ_A = θ_B is rejected when 0 ∉ [2.5th, 97.5th percentile] of "
         "θ_A^b − θ_B^b across bootstrap replicates.",
         size=10)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # DATA GENERATION PIPELINE
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "1. Data Generation Pipeline", level=1)
    para(doc,
         "The experiment uses 100 hate-speech texts with binary labels "
         "(50 no-hate / 50 hate) sampled from the UCBerkeley-DLAB Measuring Hate "
         "Speech dataset (Kennedy et al., 2020). The full audio noise pipeline "
         "consists of four stages:",
         size=10)

    stages = [
        ("Stage 1 — Text → Audio (TTS)",
         "Each text is synthesised into speech using edge-tts "
         "(voice: en-US-GuyNeural) and saved as a 16 kHz mono WAV file "
         "(100 files total)."),
        ("Stage 2 — Audio + Noise Injection",
         "Each clean WAV is mixed with one of three real-world noise recordings "
         "at 21 SNR levels (40, 30, 20, 15, 10, 5, 4, 3, 2, 1, 0, −1, −2, −3, "
         "−5, −6, −7, −8, −10, −15, −20 dB) using the power-normalised mixing "
         "formula. Noise types: white (synthetic Gaussian), babble (MUSAN speech "
         "overlay, 15 speakers), café (DEMAND CAFE real cafeteria recording). "
         "A noise-free baseline (NSR = 0) is taken directly from the original text."),
        ("Stage 3 — Audio → Text (ASR)",
         "Noisy WAVs are transcribed using faster-whisper (base model) to produce "
         "one noisy-text CSV per (noise type × SNR level), matching the format "
         "of the text-noise pipeline."),
        ("Stage 4 — LLM Classification",
         "Each noisy text is classified independently by four LLMs: "
         "GPT-3.5-turbo, GPT-5.4-nano, Gemini-2.5-Flash, Gemini-2.5-Flash-Lite. "
         "Calls are made per-item (no batching) to avoid cross-item influence. "
         "Prompt instructs binary hate-speech classification (0 = no hate, 1 = hate); "
         "tie-break: choose 0 if unsure."),
    ]

    for i, (stage_title, stage_body) in enumerate(stages, 1):
        p = doc.add_paragraph(style="List Number")
        run_t = p.add_run(stage_title + ": ")
        run_t.bold = True
        run_t.font.size = Pt(10)
        run_b = p.add_run(stage_body)
        run_b.font.size = Pt(10)

    doc.add_paragraph()
    para(doc,
         "The NSR (Noise-to-Signal power ratio) is computed as "
         "NSR = N/S = 10^(−SNR_dB / 10). "
         "The noise channel confusion probability is modelled as "
         "q(NSR) = min(α · NSR^β, 1), mirroring the text-noise formulation "
         "q(p′) = min(α · p′^β, 1).",
         size=10)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # METHODOLOGY
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "2. Methodology", level=1)
    para(doc,
         "Bootstrap resampling (B = 1,000 replicates) is used to estimate standard "
         "errors and construct 95% confidence intervals for all RI model parameters "
         "(λ per model, shared α and β). For each replicate, observations within "
         "each (model, noise-level) cell are resampled independently with replacement; "
         "the RI model is then refitted via SSE minimisation. "
         "The 95% CI is the [2.5th, 97.5th] percentile of the bootstrap distribution "
         "(percentile method).",
         size=10)

    para(doc,
         "Hypothesis tests for pairwise differences use the bootstrap CI of the "
         "difference: for two parameters θ_A and θ_B, we form D^b = θ_A^b − θ_B^b "
         "for each bootstrap replicate b and reject H₀: θ_A = θ_B at the 5% level "
         "when 0 lies outside the [2.5th, 97.5th] percentile of {D^b}. "
         "This approach is non-parametric and does not require normality of the "
         "bootstrap distribution.",
         size=10)

    # References paragraph
    p = doc.add_paragraph()
    run = p.add_run("References: ")
    run.bold = True
    run.font.size = Pt(10)
    refs = [
        "Efron, B., & Tibshirani, R. J. (1993). "
        "An Introduction to the Bootstrap. Chapman & Hall/CRC. "
        "[Bootstrap SE and percentile CI — Chapters 6, 13]",

        "Greenwood, M. C. (2022). "
        "Intermediate Statistics with R (2nd ed.). "
        "[Bootstrap CI of difference for two-sample hypothesis testing — Section 2.9]",
    ]
    for i, ref in enumerate(refs):
        sep = p.add_run("\n" if i == 0 else "\n")
        sep.font.size = Pt(10)
        r2 = p.add_run(f"[{i+1}] {ref}")
        r2.font.size = Pt(10)
        r2.italic = True

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 1 — Between-model comparison within each noise environment
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "3. Between-Model Comparison Within Each Noise Environment", level=1)
    para(doc,
         "For each noise environment we test whether pairs of LLMs show significantly "
         "different information costs (λ). A significant result (✓) indicates that "
         "0 lies outside the 95% bootstrap CI of the difference λ_A − λ_B.",
         size=10)

    # ── 1a. Text noise ────────────────────────────────────────────────────────
    heading(doc, "3a. Text Noise", level=2)

    # Bootstrap estimates table for text noise
    para(doc, "Parameter estimates (full-data fit, bootstrap SE, 95% CI):", size=10)
    lambda_rows = text_sum[text_sum["param"].str.startswith("lambda_")]
    other_rows  = text_sum[text_sum["param"].isin(["alpha", "beta"])]

    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    add_header_row(t, ["Parameter", "Estimate", "SE", "CI_lo", "CI_hi"])
    for _, r in pd.concat([lambda_rows, other_rows]).iterrows():
        add_data_row(t,
                     [r["param"].replace("lambda_", "λ "),
                      f"{r['estimate']:.4f}", f"{r['se']:.4f}",
                      f"{r['ci_lo']:.4f}", f"{r['ci_hi']:.4f}"],
                     align_center=[1, 2, 3, 4])
    set_col_widths(t, [2.2, 1.1, 1.0, 1.0, 1.0])
    doc.add_paragraph()

    # Pairwise lambda table
    para(doc, "Pairwise λ comparison (CI of difference):", size=10)
    cols = ["Model A", "Model B", "λ_A", "λ_B", "Difference", "95% CI of Diff", "Sig."]
    t2 = doc.add_table(rows=1, cols=len(cols))
    t2.style = "Table Grid"
    add_header_row(t2, cols)
    for _, r in text_pair.iterrows():
        sig = sig_bool(r["sig_ci"])
        add_data_row(t2,
                     [r["model_A"], r["model_B"],
                      f"{r['lambda_A']:.4f}", f"{r['lambda_B']:.4f}",
                      f"{r['obs_diff']:.4f}",
                      fmt_ci(r["ci_lo_diff"], r["ci_hi_diff"]),
                      SIG_MARK if sig else NOSIG_MARK],
                     sig=sig, align_center=[2, 3, 4, 5, 6])
    set_col_widths(t2, [1.5, 1.5, 0.8, 0.8, 0.9, 1.3, 0.5])
    doc.add_paragraph()

    sig_n = text_pair["sig_ci"].apply(sig_bool).sum()
    para(doc,
         f"Finding: All {len(text_pair)} model pairs show significantly different "
         f"information costs in the text noise environment ({sig_n}/{len(text_pair)} sig.). "
         f"GPT-5.2 has the smallest λ (= {text_pair.iloc[0]['lambda_A']:.4f}), "
         f"indicating the highest decision efficiency under text noise.",
         size=10, italic=True)

    doc.add_paragraph()

    # ── 1b. Audio noise ───────────────────────────────────────────────────────
    heading(doc, "3b. Audio Noise (White / Babble / Café)", level=2)
    para(doc,
         "Four LLMs were evaluated under three audio noise conditions. "
         "GPT-5.4-nano consistently shows the highest λ across all conditions, "
         "indicating the greatest information processing cost — the model is least "
         "willing to invest in acquiring better signal to improve its decision. "
         "Parameter estimates and pairwise λ comparisons are reported for each condition.",
         size=10)

    for noise in ["white", "babble", "cafe"]:
        heading(doc, NOISE_LABELS[noise], level=3)

        # Estimates
        sub_sum = aud_sum[aud_sum["noise"] == noise]
        lam_sub = sub_sum[sub_sum["param"].str.startswith("lambda_")]
        oth_sub = sub_sum[sub_sum["param"].isin(["alpha", "beta"])]

        para(doc, "Parameter estimates:", size=10)
        t = doc.add_table(rows=1, cols=5)
        t.style = "Table Grid"
        add_header_row(t, ["Parameter", "Estimate", "SE", "CI_lo", "CI_hi"])
        for _, r in pd.concat([lam_sub, oth_sub]).iterrows():
            add_data_row(t,
                         [r["param"].replace("lambda_", "λ "),
                          f"{r['estimate']:.4f}", f"{r['se']:.4f}",
                          f"{r['ci_lo']:.4f}", f"{r['ci_hi']:.4f}"],
                         align_center=[1, 2, 3, 4])
        set_col_widths(t, [2.4, 1.1, 1.0, 1.0, 1.0])
        doc.add_paragraph()

        # Pairwise
        sub_pair = aud_mod[aud_mod["noise"] == noise]
        para(doc, "Pairwise λ comparison:", size=10)
        cols = ["Model A", "Model B", "λ_A", "λ_B", "Difference", "95% CI of Diff", "Sig."]
        t2 = doc.add_table(rows=1, cols=len(cols))
        t2.style = "Table Grid"
        add_header_row(t2, cols)
        for _, r in sub_pair.iterrows():
            sig = sig_bool(r["sig_ci"])
            add_data_row(t2,
                         [r["model_A"], r["model_B"],
                          f"{r['lambda_A']:.4f}", f"{r['lambda_B']:.4f}",
                          f"{r['obs_diff']:.4f}",
                          fmt_ci(r["ci_lo_diff"], r["ci_hi_diff"]),
                          SIG_MARK if sig else NOSIG_MARK],
                         sig=sig, align_center=[2, 3, 4, 5, 6])
        set_col_widths(t2, [1.6, 1.9, 0.75, 0.75, 0.9, 1.3, 0.45])
        doc.add_paragraph()

        sig_n = sub_pair["sig_ci"].apply(sig_bool).sum()
        tot   = len(sub_pair)

        # Identify which models GPT-5.4-nano is significantly higher than
        gpt54_sig = sub_pair[
            (sub_pair["sig_ci"].apply(sig_bool)) &
            (sub_pair["model_A"].str.contains("5.4") | sub_pair["model_B"].str.contains("5.4"))
        ]
        gpt54_others = []
        for _, r in gpt54_sig.iterrows():
            other = r["model_B"] if "5.4" in r["model_A"] else r["model_A"]
            gpt54_others.append(other)

        gpt54_note = ""
        if gpt54_others:
            gpt54_note = (f" GPT-5.4-nano (λ = {sub_pair[sub_pair['model_A']=='GPT-5.4-nano'].iloc[0]['lambda_A'] if not sub_pair[sub_pair['model_A']=='GPT-5.4-nano'].empty else sub_pair[sub_pair['model_B']=='GPT-5.4-nano'].iloc[0]['lambda_B']:.4f}) "
                          f"is significantly higher than: {', '.join(gpt54_others)}.")

        para(doc,
             f"Finding ({NOISE_LABELS[noise]}): {sig_n}/{tot} model pairs are "
             f"significantly different by the CI test.{gpt54_note}",
             size=10, italic=True)
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 2 — Across-noise comparison (same model)
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "4. Across-Noise-Environment Comparison (Same Model)", level=1)
    para(doc,
         "We test whether the same model's RI parameters (λ, α, β) differ "
         "significantly across the three audio noise conditions. "
         "H₀: θ_{noise_A} = θ_{noise_B} is rejected when 0 ∉ 95% bootstrap CI "
         "of the difference.",
         size=10)

    # Lambda subtable
    heading(doc, "4a. Lambda (λ) Across Noise Types", level=2)
    lam_cross = aud_cross[aud_cross["param"].str.startswith("lambda_")]
    cols = ["Noise A", "Noise B", "Model", "λ_A", "λ_B",
            "Difference", "95% CI of Diff", "Sig."]
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Table Grid"
    add_header_row(t, cols)
    for _, r in lam_cross.iterrows():
        sig  = sig_bool(r["sig_ci"])
        model = r["param"].replace("lambda_", "")
        add_data_row(t,
                     [NOISE_LABELS.get(r["noise_A"], r["noise_A"]),
                      NOISE_LABELS.get(r["noise_B"], r["noise_B"]),
                      model,
                      f"{r['est_A']:.4f}", f"{r['est_B']:.4f}",
                      f"{r['obs_diff']:.4f}",
                      fmt_ci(r["ci_lo_diff"], r["ci_hi_diff"]),
                      SIG_MARK if sig else NOSIG_MARK],
                     sig=sig, align_center=[3, 4, 5, 6, 7])
    set_col_widths(t, [1.0, 1.0, 1.5, 0.7, 0.7, 0.85, 1.3, 0.45])
    doc.add_paragraph()

    sig_lam = lam_cross["sig_ci"].apply(sig_bool).sum()
    para(doc,
         f"Finding: {sig_lam}/{len(lam_cross)} lambda comparisons are significant. "
         f"The decision parameter λ is stable across all three audio noise types — "
         f"each model's information processing cost does not change significantly "
         f"as the type of acoustic noise changes.",
         size=10, italic=True)
    doc.add_paragraph()

    # Alpha/Beta subtable
    heading(doc, "4b. Shared Parameters (α, β) Across Noise Types", level=2)
    ab_cross = aud_cross[aud_cross["param"].isin(["alpha", "beta"])]
    cols = ["Noise A", "Noise B", "Parameter", "Value_A", "Value_B",
            "Difference", "95% CI of Diff", "Sig."]
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Table Grid"
    add_header_row(t, cols)
    for _, r in ab_cross.iterrows():
        sig = sig_bool(r["sig_ci"])
        add_data_row(t,
                     [NOISE_LABELS.get(r["noise_A"], r["noise_A"]),
                      NOISE_LABELS.get(r["noise_B"], r["noise_B"]),
                      r["param"],
                      f"{r['est_A']:.4f}", f"{r['est_B']:.4f}",
                      f"{r['obs_diff']:.4f}",
                      fmt_ci(r["ci_lo_diff"], r["ci_hi_diff"]),
                      SIG_MARK if sig else NOSIG_MARK],
                     sig=sig, align_center=[3, 4, 5, 6, 7])
    set_col_widths(t, [1.0, 1.0, 1.1, 0.75, 0.75, 0.85, 1.3, 0.45])
    doc.add_paragraph()

    sig_ab = ab_cross["sig_ci"].apply(sig_bool).sum()
    sig_rows = ab_cross[ab_cross["sig_ci"].apply(sig_bool)]
    # Build a human-readable description of significant rows
    sig_lines = []
    for _, r in sig_rows.iterrows():
        nA = NOISE_LABELS.get(r["noise_A"], r["noise_A"])
        nB = NOISE_LABELS.get(r["noise_B"], r["noise_B"])
        sig_lines.append(
            f"α differs significantly between {nA} and {nB} "
            f"(est_A={r['est_A']:.3f}, est_B={r['est_B']:.3f}, "
            f"95% CI of diff: {fmt_ci(r['ci_lo_diff'], r['ci_hi_diff'])})"
        )
    sig_desc = "; ".join(sig_lines) if sig_lines else "none"
    para(doc,
         f"Finding: {sig_ab}/{len(ab_cross)} comparisons significant. "
         f"{sig_desc}. "
         f"This suggests that the noise-sensitivity parameter α — which governs how "
         f"strongly noise degrades the signal quality — differs between white and "
         f"babble acoustic conditions, even though the decision parameter λ does not.",
         size=10, italic=True)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 3 — Summary
    # ══════════════════════════════════════════════════════════════════════════
    heading(doc, "5. Summary", level=1)

    summary_items = [
        ("Finding 1 — Within-noise model comparison (text noise)",
         f"All 6 pairwise λ comparisons are significant. "
         f"GPT-5.2 shows the smallest λ ({text_pair.iloc[0]['lambda_A']:.4f}), "
         f"indicating the highest decision efficiency; "
         f"Gemini-2.0 shows the largest ({text_pair.iloc[-1]['lambda_B']:.4f}). "
         f"All models differ significantly in information processing cost under text noise."),

        ("Finding 2 — Within-noise model comparison (audio noise)",
         "Significance is mixed across model pairs and noise types. "
         "GPT-5.4-nano consistently has the highest λ across all three audio noise conditions "
         "(white, babble, café), and is significantly higher than the other models "
         "in babble and café noise (3/3 comparisons significant), and higher than "
         "GPT-3.5-turbo in white noise. "
         "This indicates that GPT-5.4-nano incurs the greatest information processing cost "
         "and is least willing to invest in acquiring signal quality."),

        ("Finding 3 — Across-noise comparison (same model, audio)",
         f"The decision parameter λ is stable across noise types: "
         f"{sig_lam}/{len(lam_cross)} cross-noise λ comparisons are significant, "
         f"indicating that each model's information cost does not change with the type of "
         f"acoustic noise. "
         f"However, the environment parameter α differs significantly between white and "
         f"babble noise (α_white={aud_cross[(aud_cross['noise_A']=='white') & (aud_cross['noise_B']=='babble') & (aud_cross['param']=='alpha')].iloc[0]['est_A']:.3f} vs "
         f"α_babble={aud_cross[(aud_cross['noise_A']=='white') & (aud_cross['noise_B']=='babble') & (aud_cross['param']=='alpha')].iloc[0]['est_B']:.3f}), "
         f"suggesting that the noise-sensitivity of signal quality differs between these conditions."),
    ]

    for title_text, body_text in summary_items:
        p = doc.add_paragraph(style="List Bullet")
        run_title = p.add_run(title_text + ": ")
        run_title.bold = True
        run_title.font.size = Pt(10)
        run_body = p.add_run(body_text)
        run_body.font.size = Pt(10)

    # ── Save ─────────────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Report saved -> {OUT_PATH}")


if __name__ == "__main__":
    main()
